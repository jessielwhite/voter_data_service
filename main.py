from docx import Document
import csv
import pandas as pd
import re
from dotenv import load_dotenv
import os
from google.cloud import storage

load_dotenv()

project_id = storage.Client(os.get("PROJECT_ID"))
bucket = storage.Client(os.get("BUCKET_NAME"))

def authenticate_implicit_with_adc(project_id):
    buckets = project_id.list_buckets()
    print("Buckets:")
    for bucket in buckets:
        print(bucket.name)
    print("Listed all storage buckets.")

authenticate_implicit_with_adc(project_id)

def download_blob(bucket, source_voter_file_layout, destination_voter_file_layout):
    """Downloads a blob from the bucket."""

    blob = bucket.blob(source_voter_file_layout)
    voter_file = blob.download_to_filename(destination_voter_file_layout)

    print("Downloaded storage object {} from bucket {} to local file {}.".format(source_voter_file_layout, bucket, voter_file))

download_blob(bucket, source_voter_file_layout=f"{bucket}/data/Voter_File_Layout.docx", destination_voter_file_layout="Voter_File_Layout.docx")

def voter_file_layout_docx_to_csv(docx_layout, csv_layout):
    """
    Extracts text from the docx file table, removes unnecessary characters, and writes each row as a row in a CSV file.
    """
    try:
        doc = Document(docx_layout)
    except Exception as e:
        print(f"Error opening DOCX file: {e}")
        return None

    if not doc.tables:
        print(f"No tables found in '{docx_layout}'.")
        return None

    table = doc.tables[0]
    with open(csv_layout, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        for row in table.rows:
            cleaned_row = [
                re.sub(r'[^\x20-\x7E\n\t]', '', cell.text).strip()
                for cell in row.cells
            ]
            writer.writerow(cleaned_row)
    return True

voter_file_layout_docx_to_csv("Voter_File_Layout.docx", "voter_file_layout.csv")


def voter_data_txt_to_csv(txt_voter_data, csv_voter_data, delimiter=None):
    """
    Converts a delimited txt file containing county-speciifc voter data to CSV format.
    If delimiter is None, will attempt to auto-detect (comma, tab, or pipe).
    """
    # Try to auto-detect delimiter if not provided
    if delimiter is None:
        with open(txt_voter_data, 'r', encoding='utf-8') as f:
            sample = f.readline()
            if '\t' in sample:
                delimiter = '\t'
            elif '|' in sample:
                delimiter = '|'
            else:
                delimiter = ','

    with open(txt_voter_data, 'r', encoding='utf-8') as txtfile, \
         open(csv_voter_data, 'w', newline='', encoding='utf-8') as csvfile:
        reader = csv.reader(txtfile, delimiter=delimiter)
        writer = csv.writer(csvfile)
        for row in reader:
            writer.writerow(row)

voter_data_txt_to_csv('voter_data.txt', 'voter_data.csv')

def process_voter_data_csv(csv_filename):
    """
    Concatenates all values for cols beginning with "PRIMARY"/"GENERAL"/"SPECIAL" into a single val for new col VOTING_HISTORY.
    For schema normalization testing purposes.
    """
    data = pd.read_csv(csv_filename)

    cols_to_combine = [col for col in data.columns if col.startswith(('PRIMARY', 'GENERAL', 'SPECIAL'))]

    def concat_history(row):
        return ';'.join([f"{col}:{row[col]}" for col in cols_to_combine if pd.notnull(row[col]) and row[col] != ''])

    data['VOTING_HISTORY'] = data.apply(concat_history, axis=1)

    data = data.drop(columns=cols_to_combine)

    data.to_csv(csv_filename, index=False)
    return data

process_voter_data_csv('voter_data.csv')
